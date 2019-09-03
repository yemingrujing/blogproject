# -*- coding: UTF-8 -*-
from django.views.generic import View
from django.http.response import HttpResponse
from django.shortcuts import render, redirect
from django.http import StreamingHttpResponse
from blogproject import settings
import os, uuid, hashlib

from io import StringIO
from io import open
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document
# linux环境不支持win32com
# import pythoncom
# from win32com import client


# Create your views here.
class Uploads(View):
    def get(self, request):
        return render(request, 'upload/filedown.html')

    def post(self, request):
        my_file = request.FILES.get('file_name', )
        file_name = get_unique_str() + '_' + my_file.name.split('.')[-2]
        file_suffix = my_file.name.split('.')[-1]
        # 文件路径
        file_path = os.path.join(settings.UPLOAD_ROOT, file_name + '.' + file_suffix)
        f = open(file_path, 'wb')
        for i in my_file.chunks():
            f.write(i)
        f.close()
        # 转换成PDF
        if file_suffix == 'doc' or file_suffix == 'docx':
            # 文件路径
            pdf_file_path = os.path.join(settings.UPLOAD_ROOT, file_name + '.' + 'pdf')
            # word_to_doc(file_path, pdf_file_path)

        return redirect('/app/filedown')


class FileDown(View, HttpResponse):
    def get(self, request, **kwargs):
        file_path = request.GET.get("file_path")
        if file_path is None:
            # 获得父类生成的传递给模板的字典
            file_list = os.listdir(settings.UPLOAD_ROOT)
            return render(request, "upload/filedown.html",
                          context={'file_list': file_list, 'root_path': settings.UPLOAD_ROOT})
        else:
            file_name = get_unique_str() + '.' + file_path.split('.')[-1]
            resp = StreamingHttpResponse(read_file(file_path))
            resp['Content-Type'] = 'application/octet-stream'
            resp['Content-Disposition'] = 'attachment;filename="{0}"'.format(file_name)
            return resp


class FileDelete(View, HttpResponse):
    def get(self, request, **kwargs):
        file_path = request.GET.get("file_path")
        if file_path is None:
            # 获得父类生成的传递给模板的字典
            file_list = os.listdir(settings.UPLOAD_ROOT)
            return render(request, "upload/filedown.html",
                          context={'file_list': file_list, 'root_path': settings.UPLOAD_ROOT})
        else:
            if os.path.isfile(file_path):
                os.remove(file_path)
            elif os.path.isdir(file_path):
                os.rmdir(file_path)
            file_list = os.listdir(settings.UPLOAD_ROOT)
            return render(request, "upload/filedown.html",
                          context={'file_list': file_list, 'root_path': settings.UPLOAD_ROOT})


def get_unique_str():
    uuid_str = str(uuid.uuid4())
    md5 = hashlib.md5()
    md5.update(uuid_str.encode('utf-8'))
    return md5.hexdigest()


def read_file(file_name, chunk_size=512):
    with open(file_name, 'rb') as f:
        while True:
            c = f.read(chunk_size)
            if c:
                yield c
            else:
                break


def read_from_pdf(file):
    resource_manager = PDFResourceManager()
    return_str = StringIO()
    lap_params = LAParams()

    device = TextConverter(
        resource_manager, return_str, laparams=lap_params)
    process_pdf(resource_manager, device, file)
    device.close()

    content = return_str.getvalue()
    return_str.close()
    return content


def save_text_to_word(content, file_path):
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)


def pdf_to_word(file, word_file_path):
    content = read_from_pdf(file)
    save_text_to_word(content, word_file_path)


# def word_to_doc(file_path, pdf_file_path):
#     try:
#         pythoncom.CoInitialize()
#         word = client.DispatchEx("Word.Application")
#         pythoncom.CoInitialize()
#         word_doc = word.Documents.Open(file_path, ReadOnly=1)
#         word_doc.SaveAs(pdf_file_path, FileFormat=17)
#         word_doc.Close()
#     except Exception as e:
#         print(e)
#     finally:
#         # 释放资源
#         pythoncom.CoUninitialize()
